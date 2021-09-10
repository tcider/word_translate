# pip install python-docx
from typing import List, Dict
import docx
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
import PySimpleGUI as sg
# pip install PySimpleGUI
from docxtpl import DocxTemplate, RichText
import os

# pip install docxtpl


SEP = " /"
COLOR_LIST = ['BRIGHT_GREEN', 'YELLOW', 'WHITE', 'BLACK', 'BLUE', 'DARK_BLUE', 'DARK_RED', 'DARK_YELLOW', 'GRAY_25',
              'GRAY_50', 'GREEN', 'PINK', 'RED', 'TEAL', 'TURQUOISE', 'VIOLET']
COLOR = 'BRIGHT_GREEN'
ENG_LETTERS_IN_TEXT = 0.35
ENG_ALPHABET = set('abcdefghijklmnopqrstuvwxyz')
RUS_ALPHABET = set('абвгдеёжзийклмнопрстуфхцчшщъыьэюя')
NUMS = set('0123456789')


def is_rus(text):
    text = str(text)
    if not text:
        return False
    return not RUS_ALPHABET.isdisjoint(text.lower())


def is_eng(text):
    text = str(text)
    if not text:
        return False
    return not ENG_ALPHABET.isdisjoint(text.lower())


def is_eng_big(text):
    text = str(text)
    if not text:
        return False
    i = 0
    j = 0
    for c in text.lower():
        if c in ENG_ALPHABET:
            i += 1
        elif c in RUS_ALPHABET:
            j += 1
    if i / (i + j) >= ENG_LETTERS_IN_TEXT:
        return True
    return False


def is_empty(text):
    text = str(text)
    if not text:
        return True
    return str.isspace(text)


def is_num(text):
    text = str(text)
    if not text:
        return False
    return not NUMS.isdisjoint(text)


def tpl(i: int) -> str:
    return "{{_" + str(i) + "}}"


def un_tpl(text: str):
    return text.replace("{{", "").replace("}}", "")


def main_window(template_file=None, translate_file=None):  # TODO проверки на не те и пустые файлы
    global COLOR
    layout = [
        #[sg.Text('Настройки', font=("Helvetica", 20))],
        [sg.Text('Разделитель'), sg.InputText(SEP, size=(2, 1), key='sep'), sg.Text('     Цвет выделения перевода'),
         sg.Drop(values=(COLOR_LIST), default_value=COLOR, key='color', size=(15, 1))],
        [sg.Text('_' * 86)],
        [sg.Text('Индивидуальная обработка', font=("Helvetica", 20))],
        [sg.Text('1 шаг: Выберите исходный файл .docx')],
        [sg.Text('Исходный файл            '), sg.InputText(size=(55, 1), key='source'), sg.FileBrowse('Обзор')],
        [sg.Text('2 шаг: Выберите шаблон   _template.docx и перевод   _translate.docx')],
        [sg.Text('Шаблон _template.docx '), sg.InputText(template_file, size=(55, 1), key='template'),
         sg.FileBrowse('Обзор')],
        [sg.Text('Перевод _translate.docx'), sg.InputText(translate_file, size=(55, 1), key='translate'),
         sg.FileBrowse('Обзор')],
        [sg.Submit('Подготовить шаблон'), sg.Text(), sg.Submit('Вставить перевод')],
        [sg.Text('_' * 86)],
        [sg.Text('Групповая автоматическая обработка', font=("Helvetica", 20))],
        [sg.Text('Выберите папку с файлами .docx'), sg.InputText(size=(47, 1), key='dir'), sg.FolderBrowse('Обзор')],
        [sg.Text('1 шаг:'), sg.Submit('Подготовить шаблоны'), sg.Text('    2шаг: После перевода в _translate.docx '),
         sg.Submit('Вставить переводы')],
        [sg.Output(size=(83, 25), key='-OUTPUT-')],
        #[sg.Text('_' * 86)],
        #[sg.Cancel('Выход')]
    ]

    window = sg.Window('Подготовка .docx файла для перевода', layout)
    while True:
        event, values = window.read()
        if event in (None, 'Выход', sg.WIN_CLOSED):
            window.close()
            return None
        if event == 'Подготовить шаблон':
            window['-OUTPUT-'].update('')
            if values['color']:
                COLOR = values['color']
            template_file, translate_file = make_template_file2(values['source'])
            if not template_file:
                sg.popup_error("Неверный файл, либо открыт файл ' _template.docx'")
                window.close()
                main_window()
            #translate_file = make_translate_file(translate_dict, template_file)
            if not translate_file:
                sg.popup_error("Закройте файл ' _translate.docx'")
                window.close()
                main_window()
            sg.popup(f"Шаблон:\n{template_file}\nи Файл для перевода:\n{translate_file}\nГотовы")
            window.close()
            main_window(template_file, translate_file)
        if event == 'Вставить перевод':
            window['-OUTPUT-'].update('')
            result_file = make_result_file(values['template'], values['translate'])
            if not result_file:
                sg.popup_error("Неверные файлы, либо открыт ' _resulte.docx'")
                window.close()
                main_window(values['template'], values['translate'])
            sg.popup(f"Перевод сохранен в файл:\n{result_file}")
        if event == 'Подготовить шаблоны':
            if not values['dir']:
                sg.popup_error("Не выбрана директория!")
                window.close()
                main_window()
            window['-OUTPUT-'].update('')
            for file in os.listdir(values['dir']):
                if file.split('.')[-1] == "docx" and file.find("_template") == -1 and file.find("_translate") == -1 and file.find("_result") == -1 and file.find("~$") == -1:
                    file_name = values['dir'] + "/" + file
                    template_file, translate_file = make_template_file2(file_name)
                    if not template_file:
                        print(f"Закройте файл {file} _template. Шаблон не подготовлен!")
                    if not translate_file:
                        print(f"Закройте файл {file} _translate. Шаблон не подготовлен!")
                    if template_file and translate_file:
                        print(f"Шаблон:\n{template_file}\nи Файл для перевода:\n{translate_file}\nГотовы")
            sg.popup("Все файлы обработаны")
        if event == 'Вставить переводы':
            if not values['dir']:
                sg.popup_error("Не выбрана директория!")
                window.close()
                main_window()
            window['-OUTPUT-'].update('')
            for file in os.listdir(values['dir']):
                if file.split('.')[-1] == "docx" and file.find("_template") != -1:
                    template_file = values['dir'] + "/" + file
                    translate_file = template_file.replace("_template", "_translate")
                    print(f"Старт обработки шаблона {file}")
                    result_file = make_result_file(template_file, translate_file)
                    if not result_file:
                        print(f"Закройте файл {file} _result. Шаблон не подготовлен!")
                    if result_file:
                        print(f"Перевод сохранен в файл:\n{result_file}")
            sg.popup("Все файлы обработаны")


def process_paragraph(paragraph, res, i): #FIXME OLD version to del
    if is_empty(paragraph.text):
        return res, i

    paragraph_style = paragraph.style
    run_style = [paragraph.runs[0].bold,
                 paragraph.runs[0].underline,
                 paragraph.runs[0].italic,
                 paragraph.runs[0].font.name,
                 paragraph.runs[0].font.size,
                 paragraph.runs[0].style]
    lines = paragraph.text.split('\n')
    tmp = i
    for j in range(len(lines)):
        if is_empty(lines[j]):
            continue
        if lines[j].find(SEP) == -1:
            if is_rus(lines[j]):
                res[i] = lines[j]
                lines[j] += SEP + " " + tpl(i)
                i += 1
        else:
            paragraph_list = lines[j].split(SEP)
            if is_empty(paragraph_list[-2]):
                continue
            if is_rus(paragraph_list[-2]) and is_empty(paragraph_list[-1]):
                if j == len(lines) - 1 or is_rus(lines[j + 1].split(SEP)[0]):
                    res[i] = paragraph_list[-2]
                    lines[j] += " " + tpl(i)
                    i += 1
            elif is_rus(paragraph_list[-2]) and not is_eng(paragraph_list[-1]):
                res[i] = paragraph_list[-2]
                paragraph_list[-1] = " " + tpl(i) + paragraph_list[-1]
                i += 1
                lines[j] = SEP.join(paragraph_list)
    if i > tmp:
        paragraph.style = paragraph_style

        paragraph.text = ''
        text = '\n'.join(lines)
        text_list = text.split(SEP + " ")
        full_text_list = []
        for text in text_list:
            full_text_list2 = []
            text_list2 = text.split('\n')
            for text2 in text_list2:
                full_text_list2.append(text2)
                full_text_list2.append('\n')
            full_text_list2.pop(-1)
            full_text_list.extend(full_text_list2)
            full_text_list.append(SEP + " ")
        full_text_list.pop(-1)
        for text in full_text_list:
            run = paragraph.add_run(text, style=run_style[5])
            run.bold = run_style[0]
            run.underline = run_style[1]
            run.italic = run_style[2]
            run.font.name = run_style[3]
            run.font.size = run_style[4]
            if text.find("{{") != -1:
                run.font.highlight_color = eval(f"WD_COLOR_INDEX.{COLOR}")
    return res, i


def process_paragraph2(paragraph, tr_table, i): #FIXME OLD version to del
    if is_empty(paragraph.text):
        return i

    # if paragraph.text.find("Запись в журнале ПНР.") != -1:
    #     while True:
    #         print(paragraph.text)

    paragraph_style = paragraph.style
    run_style = [paragraph.runs[0].bold,
                 paragraph.runs[0].underline,
                 paragraph.runs[0].italic,
                 paragraph.runs[0].font.name,
                 paragraph.runs[0].font.size,
                 paragraph.runs[0].style]
    lines = paragraph.text.split('\n')
    tmp = i
    for j in range(len(lines)):
        if is_empty(lines[j]):
            continue

        if lines[j].find(SEP) == -1 and lines[j].rstrip()[-1] == SEP[1]:
            lines[j] = lines[j].replace(SEP[1], SEP)

        if lines[j].find(SEP) == -1: #
            if is_rus(lines[j]):
                cells = tr_table.add_row().cells
                cell0 = cells[0]
                cell1 = cells[1]
                cell0.text = tpl(i)
                cell1.text = lines[j]
                print(cell0.text, cell1.text) #FIXME del
                lines[j] += SEP + " " + tpl(i)
                i += 1
        else:
            paragraph_list = lines[j].split(SEP)
            if is_empty(paragraph_list[-2]):
                continue
            if is_rus(paragraph_list[-2]) and is_empty(paragraph_list[-1]):
                if j == len(lines) - 1 or is_empty(lines[j + 1]) or lines[j + 1].split(SEP)[0]:
                    cells = tr_table.add_row().cells
                    cell0 = cells[0]
                    cell1 = cells[1]
                    cell0.text = tpl(i)
                    cell1.text = paragraph_list[-2]
                    print(cell0.text, cell1.text)  # FIXME del
                    lines[j] += " " + tpl(i)
                    i += 1
            elif is_rus(paragraph_list[-2]) and not is_eng(paragraph_list[-1]):
                cells = tr_table.add_row().cells
                cell0 = cells[0]
                cell1 = cells[1]
                cell0.text = tpl(i)
                cell1.text = paragraph_list[-2]
                print(cell0.text, cell1.text)  # FIXME del
                paragraph_list[-1] = " " + tpl(i) + paragraph_list[-1]
                i += 1
                lines[j] = SEP.join(paragraph_list)
    if i > tmp:
        paragraph.style = paragraph_style

        paragraph.text = ''
        text = '\n'.join(lines)
        text_list = text.split(SEP + " ")
        full_text_list = []
        for text in text_list:
            full_text_list2 = []
            text_list2 = text.split('\n')
            for text2 in text_list2:
                full_text_list2.append(text2)
                full_text_list2.append('\n')
            full_text_list2.pop(-1)
            full_text_list.extend(full_text_list2)
            full_text_list.append(SEP + " ")
        full_text_list.pop(-1)
        for text in full_text_list:
            run = paragraph.add_run(text, style=run_style[5])
            run.bold = run_style[0]
            run.underline = run_style[1]
            run.italic = run_style[2]
            run.font.name = run_style[3]
            run.font.size = run_style[4]
            if text.find("{{") != -1:
                run.font.highlight_color = eval(f"WD_COLOR_INDEX.{COLOR}")
    return i


#FIXME OLD vesrion to del
def make_template_file(file: str):
    i = 0
    res = {}
    try:
        doc = docx.Document(file)
    except:
        return (res, None)

    ##### Обработка текста(не в таблицах) #####
    for paragraph in doc.paragraphs:
        if is_empty(paragraph.text):
            continue
        res, i = process_paragraph(paragraph, res, i)

    ################# Tables #####################

    for table in doc.tables:
        """
        TODO
        1.docx - page38-43 не прописываются нек-е ячейки
        2.docx - п10 не там перевод и п7 пропущены нек-е подпункты
        """
        if table.style.name == 'Normal Table':  # Nromal table Если нет стиля
            for row in table.rows:
                if not len(row.cells):
                    continue
                else:
                    for cell in row.cells:
                        if is_rus(cell.text) and cell.text.find("{{_") == -1 and (not is_eng_big(cell.text) or cell.text.strip()[-1] == SEP[-1]): #FIXME 16.08
                            for paragraph in cell.paragraphs:
                                res, i = process_paragraph(paragraph, res, i)
        else:
            tmp_loc = None
            for row in table.rows:
                if not len(row.cells):
                    continue
                else:
                    rus_text = None
                    for ic, cell in enumerate(row.cells):
                        # print("cell", cell.text) #FIXME del
                        tc = cell._tc
                        cell_loc = None
                        try:
                            cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
                        except:
                            pass
                        else:
                            if cell_loc == tmp_loc and cell.text.find("{{_") == -1 and is_rus(cell.text) and ic == len(row.cells) - 1 and (not is_eng_big(cell.text) or cell.text.strip()[-1] == SEP[-1]):  # and ic == 1
                                for paragraph in cell.paragraphs:
                                    res, i = process_paragraph(paragraph, res, i)
                                    # print(paragraph.text) #FIXME del
                            if cell_loc == tmp_loc:
                                continue
                        tmp_loc = cell_loc
                        if rus_text and not is_eng(cell.text) and (
                                is_num(cell.text) or is_empty(cell.text) or not len(cell.text)):  # is_rus(cell.text) :
                            paragraph = cell.paragraphs[0]
                            # paragraph.style = paragraph_style
                            # print(cell.text)
                            # run = paragraph.add_run(tpl(i), style=run_style[5])
                            run = paragraph.add_run(tpl(i))
                            # paragraph.text = tpl(i)
                            # run = paragraph.runs[0]
                            # cell.text = tpl(i)
                            res[i] = rus_text
                            i += 1
                            run.bold = run_style[0]
                            run.underline = run_style[1]
                            run.italic = run_style[2]
                            run.font.name = run_style[3]
                            run.font.size = run_style[4]
                            run.font.highlight_color = eval(f"WD_COLOR_INDEX.{COLOR}")

                        if is_rus(cell.text) and cell.text.find("{{_") == -1 and (not is_eng_big(cell.text) or cell.text.strip()[-1] == SEP[-1]):
                            if cell.text.find(SEP) != -1:
                                for paragraph in cell.paragraphs:
                                    res, i = process_paragraph(paragraph, res, i)
                            else:
                                rus_text = cell.text
                                paragraph = cell.paragraphs[0]
                                paragraph_style = paragraph.style
                                run_style = [paragraph.runs[0].bold,
                                             paragraph.runs[0].underline,
                                             paragraph.runs[0].italic,
                                             paragraph.runs[0].font.name,
                                             paragraph.runs[0].font.size,
                                             paragraph.runs[0].style]
                        else:
                            rus_text = None

            #         print('-' * 20)
            #     print('=' * 20)
            # print('+' * 20)

    ###############################################

    template_file = file.split(".docx")[-2] + "_template.docx"
    try:
        doc.save(template_file)
    except:
        template_file = None
    return res, template_file


def make_template_file2(file: str):
    i = 0
    try:
        doc = docx.Document(file)
    except:
        return None
    tr_doc = docx.Document()
    tr_table = tr_doc.add_table(rows=1, cols=3)
    tr_table.style = 'Table Grid'
    cells = tr_table.rows[0].cells
    cell0 = cells[0]
    cell0.text = 'Key'
    cell0.width = Inches(0.8)
    cell1 = cells[1]
    cell1.text = 'Russian'
    cell1.width = Inches(3)
    cell2 = cells[2]
    cell2.text = 'Translate'
    cell2.width = Inches(3)

    ##### Обработка текста(не в таблицах) #####
    for paragraph in doc.paragraphs:
        if is_empty(paragraph.text):
            continue
        i = process_paragraph2(paragraph, tr_table, i)

    ################# Tables #####################

    for table in doc.tables:
        """
        TODO
        1.docx - page38-43 не прописываются нек-е ячейки
        2.docx - п10 не там перевод и п7 пропущены нек-е подпункты
        """
        flag = 0
        # if table.style.name == 'Normal Table':  # Nromal table Если нет стиля
        if False: # Отключил проверку и деление таблиц по стилям
            for row in table.rows:
                try:
                    if not len(row.cells):
                        continue
                except:
                    continue
                for cell in row.cells:
                    if is_rus(cell.text) and cell.text.find("{{_") == -1 and (not is_eng_big(cell.text) or cell.text.find(SEP) != -1):  # cell.text.rstrip(" \n\r\t")[-1] == SEP[-1]
                        for paragraph in cell.paragraphs:
                            i = process_paragraph2(paragraph, tr_table, i)
        else:
            tmp_loc = None
            for row in table.rows:
                try:
                    len_row_cell = len(row.cells)
                except:
                    continue
                if not len_row_cell:
                    continue
                else:
                    rus_text = None
                    for ic, cell in enumerate(row.cells):
                        # print("cell", cell.text) #FIXME del
                        if cell.text.find("{{_") == -1 and is_rus(cell.text) and ((not is_eng_big(cell.text) and cell.text.find(SEP) != -1) or cell.text.rstrip()[-1] == SEP[1]):
                            flag += 1
                            for paragraph in cell.paragraphs:
                                i = process_paragraph2(paragraph, tr_table, i)
                            continue

                        # tc = cell._tc
                        # cell_loc = None
                        # try:
                        #     cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
                        # except:
                        #     pass
                        # else:
                        #     if cell_loc == tmp_loc and cell.text.find("{{_") == -1 and is_rus(cell.text) and ic == len(row.cells) - 1 and (not is_eng_big(cell.text) or cell.text.strip()[-1] == SEP[-1]):
                        #     #if cell_loc == tmp_loc and cell.text.find("{{_") == -1 and is_rus(cell.text) and not is_eng_big(cell.text) and cell.text.strip()[-1] == SEP[-1]:
                        #         for paragraph in cell.paragraphs:
                        #             i = process_paragraph2(paragraph, tr_table, i)
                        #             # print(paragraph.text) #FIXME del
                        #     if cell_loc == tmp_loc:
                        #         continue
                        # tmp_loc = cell_loc
                        " /  "
                        " /  "

                        if rus_text and cell.text.find("{{_") == -1 and not is_eng(cell.text) and (is_num(cell.text) or is_empty(cell.text) or len(cell.text) == 1):  # is_rus(cell.text) or not len(cell.text)
                            paragraph = cell.paragraphs[0]
                            # paragraph.style = paragraph_style
                            # print(cell.text)
                            # run = paragraph.add_run(tpl(i), style=run_style[5])
                            run = paragraph.add_run(tpl(i))
                            # paragraph.text = tpl(i)
                            # run = paragraph.runs[0]
                            # cell.text = tpl(i)
                            cells = tr_table.add_row().cells
                            cell0 = cells[0]
                            cell1 = cells[1]
                            cell0.text = tpl(i)
                            cell1.text = rus_text
                            print(cell0.text, cell1.text)  # FIXME del
                            i += 1
                            run.bold = run_style[0]
                            run.underline = run_style[1]
                            run.italic = run_style[2]
                            run.font.name = run_style[3]
                            run.font.size = run_style[4]
                            run.font.highlight_color = eval(f"WD_COLOR_INDEX.{COLOR}")
                        # elif rus_text and cell.text.find("{{_") == -1 and is_eng(cell.text):
                        #     rus_text = None

                        if is_rus(cell.text) and cell.text.find("{{_") == -1 and not is_eng_big(cell.text):
                            if cell.text.find(SEP) != -1 or flag > 2:
                                flag += 1
                                for paragraph in cell.paragraphs:
                                    i = process_paragraph2(paragraph, tr_table, i)
                            else:
                                rus_text = cell.text
                                paragraph = cell.paragraphs[0]
                                #paragraph_style = paragraph.style
                                if paragraph.runs:
                                    run_style = [paragraph.runs[0].bold,
                                                 paragraph.runs[0].underline,
                                                 paragraph.runs[0].italic,
                                                 paragraph.runs[0].font.name,
                                                 paragraph.runs[0].font.size,
                                                 paragraph.runs[0].style]

                        else:
                            rus_text = None

    ###############################################

    template_file = file.split(".docx")[-2] + "_template.docx"
    translate_file = template_file.replace("_template", "_translate")
    try:
        doc.save(template_file)
    except:
        template_file = None
    try:
        tr_doc.save(translate_file)
    except:
        translate_file = None
    return template_file, translate_file


def make_translate_file(res: Dict, template_file: str) -> str:
    translate_file = template_file.replace("_template", "_translate")
    doc = docx.Document()
    table = doc.add_table(rows=len(res), cols=3)
    # table.columns[0].width = 1097280
    table.style = 'Table Grid'
    row = 0
    for key in res:
        cells = table.add_row().cells
        # cell0 = table.cell(row, 0)
        cell0 = cells[0]
        cell0.text = tpl(key)
        cell0.width = Inches(0.8)  # 1097280
        # cell1 = table.cell(row, 1)
        cell1 = cells[1]
        cell1.text = res[key]
        cell1.width = Inches(3)
        cell2 = cells[2]
        # table.cell(row, 2).width = Inches(3)
        cell2.width = Inches(3)
        row += 1
        print(row) #FIXME 16.08
    try:
        doc.save(translate_file)
    except:
        translate_file = None
    return translate_file


def make_result_file(template_file, translate_file):
    result_file = template_file.replace("_template", "_result")
    res = {}
    try:
        doc = docx.Document(translate_file)
    except:
        return None
    for row in doc.tables[0].rows:
        index = un_tpl(row.cells[0].text)
        # rt_text = RichText()
        # style = row.cells[0].paragraphs[0].style
        # rt_text.add(row.cells[2].text, style=style, color='#0000ff')
        res[index] = row.cells[2].text.replace(SEP, '')
        print(index, res[index])
    print(f"Идет вставка перевода в шаблон {template_file}")
    try:
        template = DocxTemplate(template_file)
    except:
        return None
    template.render(res)
    try:
        template.save(result_file)
    except:
        result_file = None
    return result_file


def main():
    main_window()

    # template_file, translate_file = make_template_file2("test7/1.docx")





    # doc = docx.Document("test3/2.docx")
    # for it in range(len(doc.tables)):
    #     table = doc.tables[it]
    #     for ir in range(len(table.rows)):
    #         for ic in range(len(table.columns)):
    #             print(it, ir, ic)


if __name__ == "__main__":
    main()
